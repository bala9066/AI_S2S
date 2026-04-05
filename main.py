"""
Hardware Pipeline — FastAPI backend.

Design principles applied here:
- Thin route handlers: parse request → call service → return response.
- No business logic in this file (lives in services/).
- CORS restricted to known origins only.
- Secrets validated at startup — missing keys cause an early, clear error.
- Pipeline runs as BackgroundTask (non-blocking, UI polls for status).
- Structured logging via Python logging throughout.
"""

import functools
import logging
import os
import threading
from contextlib import asynccontextmanager
from datetime import datetime
from typing import Optional

from fastapi import BackgroundTasks, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse, RedirectResponse
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel
from typing import List

from config import settings
from logging_config import configure_logging

configure_logging()
log = logging.getLogger("hardware_pipeline.api")


# ── Startup / shutdown ────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    log.info("startup.begin", extra={"env": settings.app_env})

    # 1. Validate secrets — fail fast with a clear message
    _validate_secrets()

    # 2. Initialise DB (creates tables if they don't exist)
    from database.models import get_engine
    get_engine()
    log.info("startup.db_ready")

    # 3. Ensure output directory exists
    settings.output_dir.mkdir(parents=True, exist_ok=True)

    # 4. Seed ChromaDB component index in background — non-blocking
    def _seed_chroma():
        try:
            from tools.seed_components import seed_if_empty
            seed_if_empty()
            log.info("startup.chroma_ready")
        except Exception as exc:
            log.debug("startup.chroma_seed_skipped: %s (optional)", exc)

    threading.Thread(target=_seed_chroma, daemon=True, name="chroma-seed").start()

    log.info("startup.complete", extra={"air_gapped": settings.is_air_gapped})
    yield
    log.info("shutdown.complete")


def _validate_secrets() -> None:
    """
    Fail fast if required secrets are missing.
    Logs a clear warning for optional keys so operators know what's degraded.
    """
    if not settings.has_any_llm_key:
        if settings.app_env == "production":
            raise RuntimeError(
                "No LLM API key configured (ANTHROPIC_API_KEY or GLM_API_KEY). "
                "Set at least one before starting in production."
            )
        log.warning("startup.no_llm_key — running in air-gap/Ollama mode")

    optional_keys = {
        "DIGIKEY_CLIENT_ID": settings.digikey_client_id,
        "MOUSER_API_KEY": settings.mouser_api_key,
        "OPENAI_API_KEY": settings.openai_api_key,
    }
    for name, val in optional_keys.items():
        if not val:
            log.info("startup.optional_key_missing: %s (degraded mode)", name)


# ── App factory ────────────────────────────────────────────────────────────────

app = FastAPI(
    title=settings.app_name,
    description="AI-powered hardware design automation pipeline",
    version="2.0.0",
    lifespan=lifespan,
    docs_url="/docs" if settings.debug else None,   # hide Swagger in prod
    redoc_url=None,
)

# CORS: restrict to known origins only (never wildcard in any real deploy)
_ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Content-Type", "Authorization"],
)

# ── Password Gate (optional) ───────────────────────────────────────────────────
# Set APP_PASSWORD env var to enable. Leave empty to disable (open access).
# Uses a signed cookie — no database, no sessions library needed.
# Protects all routes except /health and /login.

_APP_PASSWORD = settings.app_password or os.environ.get("APP_PASSWORD", "")
_COOKIE_NAME = "hp_auth"
_COOKIE_MAX_AGE = 60 * 60 * 24 * 7  # 7 days

def _make_token(password: str) -> str:
    """Simple HMAC token so the cookie can't be forged without knowing the password."""
    import hmac
    import hashlib
    return hmac.new(password.encode(), b"hardware-pipeline-auth", hashlib.sha256).hexdigest()

_LOGIN_PAGE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Hardware Pipeline — Login</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@700;800&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    background: #070b14;
    color: #e2e8f0;
    font-family: 'DM Mono', monospace;
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    background-image: radial-gradient(circle at 50% 50%, rgba(0,198,167,0.04) 0%, transparent 70%);
  }
  .card {
    background: #1a2235;
    border: 1px solid rgba(0,198,167,0.25);
    border-radius: 12px;
    padding: 48px 44px;
    width: 100%;
    max-width: 400px;
    box-shadow: 0 0 40px rgba(0,198,167,0.08);
    text-align: center;
  }
  .logo { font-family: 'Syne', sans-serif; font-size: 26px; font-weight: 800; margin-bottom: 4px; }
  .logo span { color: #00c6a7; }
  .sub { font-size: 10px; color: #00c6a7; letter-spacing: 0.15em; margin-bottom: 32px; }
  label { display: block; font-size: 11px; color: #64748b; letter-spacing: 0.08em; margin-bottom: 8px; text-align: left; }
  input[type=password] {
    width: 100%; padding: 12px 14px;
    background: #0d1220; border: 1px solid rgba(42,58,80,0.8);
    border-radius: 6px; color: #e2e8f0; font-family: 'DM Mono', monospace;
    font-size: 14px; outline: none; margin-bottom: 18px;
    transition: border-color 0.2s;
  }
  input[type=password]:focus { border-color: #00c6a7; }
  button {
    width: 100%; padding: 12px;
    background: #00c6a7; border: none; border-radius: 6px;
    color: #070b14; font-family: 'Syne', sans-serif;
    font-size: 14px; font-weight: 700; cursor: pointer;
    letter-spacing: 0.05em; transition: opacity 0.2s;
  }
  button:hover { opacity: 0.88; }
  .err { color: #ef4444; font-size: 12px; margin-bottom: 14px; }
</style>
</head>
<body>
<div class="card">
  <div class="logo">Hardware <span>Pipeline</span></div>
  <div class="sub">DATA PATTERNS · CODE KNIGHTS</div>
  <form method="POST" action="/login">
    <label>ACCESS PASSWORD</label>
    <input type="password" name="password" placeholder="Enter password" autofocus>
    {error}
    <button type="submit">ENTER →</button>
  </form>
</div>
</body>
</html>"""

class PasswordGateMiddleware(BaseHTTPMiddleware):
    """Block all routes behind a password if APP_PASSWORD is set."""

    # Routes that bypass the gate entirely
    _OPEN = {"/health", "/login"}

    async def dispatch(self, request: Request, call_next):
        if not _APP_PASSWORD:
            return await call_next(request)  # gate disabled

        path = request.url.path
        if path in self._OPEN or path.startswith("/login"):
            return await call_next(request)

        # Check cookie
        token = request.cookies.get(_COOKIE_NAME, "")
        if token == _make_token(_APP_PASSWORD):
            return await call_next(request)

        # Not authenticated — redirect to login
        return RedirectResponse(url=f"/login?next={path}", status_code=302)

if _APP_PASSWORD:
    app.add_middleware(PasswordGateMiddleware)
    log.info("password_gate.enabled")


@app.get("/login", response_class=HTMLResponse, tags=["ops"])
async def login_page(next: str = "/app"):
    return HTMLResponse(_LOGIN_PAGE.replace("{error}", ""))


@app.post("/login", tags=["ops"])
async def login_submit(request: Request, next: str = "/app"):
    form = await request.form()
    password = form.get("password", "")
    if password == _APP_PASSWORD:
        response = RedirectResponse(url=next, status_code=302)
        response.set_cookie(
            key=_COOKIE_NAME,
            value=_make_token(_APP_PASSWORD),
            max_age=_COOKIE_MAX_AGE,
            httponly=True,
            samesite="lax",
        )
        return response
    error_html = '<div class="err">Incorrect password. Try again.</div>'
    return HTMLResponse(_LOGIN_PAGE.replace("{error}", error_html), status_code=401)


# ── Service singletons (created once per process) ─────────────────────────────
# lru_cache ensures a single instance is reused for the lifetime of the process.

@functools.lru_cache(maxsize=1)
def _project_svc():
    from services.project_service import ProjectService
    return ProjectService()

@functools.lru_cache(maxsize=1)
def _chat_svc():
    from services.chat_service import ChatService
    return ChatService()

@functools.lru_cache(maxsize=1)
def _pipeline_svc():
    from services.pipeline_service import PipelineService
    return PipelineService()


# ── Health ─────────────────────────────────────────────────────────────────────

@app.get("/health", tags=["ops"])
async def health_check():
    return {
        "status": "healthy",
        "app": settings.app_name,
        "environment": settings.app_env,
        "air_gapped": settings.is_air_gapped,
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }


# ── Projects ───────────────────────────────────────────────────────────────────

@app.post("/api/v1/projects", status_code=201, tags=["projects"])
async def create_project(body: dict):
    name = (body.get("name") or "").strip()
    if not name:
        raise HTTPException(400, "name is required")
    try:
        return _project_svc().create(
            name=name,
            description=body.get("description", ""),
            design_type=body.get("design_type", "rf"),
        )
    except Exception as exc:
        log.exception("api.create_project_failed")
        raise HTTPException(500, str(exc))


@app.get("/api/v1/projects", tags=["projects"])
async def list_projects():
    return _project_svc().list_all()


@app.get("/api/v1/projects/{project_id}", tags=["projects"])
async def get_project(project_id: int):
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    return proj


def _resolve_output_dir(proj: dict) -> Optional[str]:
    """
    Return the output directory for a project, with fallback derivation.

    Priority:
    1. DB-stored output_dir (absolute or relative path that exists on disk)
    2. Derived from project name using the same StorageAdapter logic
       (handles projects created before output_dir was reliably written, or
        where the DB column was left empty due to a failed project creation)
    """
    stored = (proj.get("output_dir") or "").strip()
    if stored and os.path.isdir(stored):
        return stored

    # Fallback: derive from project name using StorageAdapter.project_dir logic
    name = (proj.get("name") or "").strip()
    if name:
        safe = name.replace(" ", "_").lower()
        # Try relative (server started from project root) and absolute via settings
        candidates = [
            os.path.join("output", safe),
            str(settings.output_dir / safe),
        ]
        for candidate in candidates:
            if os.path.isdir(candidate):
                log.info(
                    "documents.output_dir_derived",
                    extra={"project_id": proj.get("id"), "derived": candidate, "stored": stored or "(empty)"},
                )
                return candidate

    log.warning(
        "documents.output_dir_missing",
        extra={"project_id": proj.get("id"), "stored": stored or "(empty)", "name": proj.get("name")},
    )
    return None


@app.get("/api/v1/projects/{project_id}/documents/{filename:path}", tags=["projects"])
async def get_document(project_id: int, filename: str):
    # :path type captures slashes, so qt_gui/ControlPanel.cpp works as-is
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "Project output directory not found — run Phase 1 first")

    # Guard against path traversal
    base = os.path.realpath(output_dir)
    file_path = os.path.realpath(os.path.join(output_dir, filename))
    if not file_path.startswith(base):
        raise HTTPException(400, "Invalid filename")
    if not os.path.exists(file_path):
        raise HTTPException(404, f"Document {filename} not found")

    return FileResponse(file_path)


@app.get("/api/v1/projects/{project_id}/documents", tags=["projects"])
async def list_documents(project_id: int):
    """List all available output files for a project (flat + one level deep for qt_gui etc.)."""
    proj = _project_svc().get(project_id)
    if not proj:
        return []

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        return []

    try:
        files = []
        for entry in os.scandir(output_dir):
            if entry.is_file():
                files.append({"name": entry.name, "size": entry.stat().st_size})
            elif entry.is_dir():
                # Include one level of subdirectory files (e.g. qt_gui/, .github/workflows/)
                try:
                    for sub in os.scandir(entry.path):
                        if sub.is_file():
                            rel = f"{entry.name}/{sub.name}"
                            files.append({"name": rel, "size": sub.stat().st_size})
                except OSError:
                    pass
        return sorted(files, key=lambda x: x["name"])
    except OSError as exc:
        log.warning("documents.list_failed", extra={"project_id": project_id, "error": str(exc)})
        return []


# ── Chat (Phase 1) ─────────────────────────────────────────────────────────────

@app.post("/api/v1/projects/{project_id}/chat", tags=["chat"])
async def chat(project_id: int, body: dict):
    """Send a message to the Phase 1 requirements agent."""
    message = (body.get("message") or "").strip()
    if not message:
        raise HTTPException(400, "message is required")

    try:
        result = await _chat_svc().send_message(project_id, message)
        log.info("api.chat_ok",
                 extra={"project_id": project_id, "phase_complete": result.get("phase_complete")})
        return result
    except ValueError as exc:
        raise HTTPException(404, str(exc))
    except Exception as exc:
        log.exception("api.chat_failed", extra={"project_id": project_id})
        raise HTTPException(500, str(exc))


# ── Pipeline (P2→P8c background execution) ────────────────────────────────────

@app.post("/api/v1/projects/{project_id}/pipeline/run", tags=["pipeline"])
async def run_pipeline(project_id: int, background_tasks: BackgroundTasks):
    """
    Start the full P2→P8c pipeline as a background task.
    Returns immediately; UI should poll GET /projects/{id} for status.
    """
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    if _project_svc().get_phase_status(project_id, "P1") != "completed":
        raise HTTPException(400, "Phase 1 must be completed before running the pipeline")

    svc = _pipeline_svc()
    background_tasks.add_task(svc.run_pipeline, project_id)
    log.info("api.pipeline_started", extra={"project_id": project_id})
    return {"status": "pipeline_started", "project_id": project_id}


VALID_PHASES = {"P1", "P2", "P3", "P4", "P5", "P6", "P7a", "P8a", "P8b", "P8c"}

@app.post("/api/v1/projects/{project_id}/phases/{phase_id}/execute", tags=["pipeline"])
async def execute_single_phase(project_id: int, phase_id: str, background_tasks: BackgroundTasks):
    """Execute one specific phase as a background task."""
    if phase_id not in VALID_PHASES:
        raise HTTPException(400, f"Invalid phase '{phase_id}'. Must be one of: {sorted(VALID_PHASES)}")
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    try:
        svc = _pipeline_svc()
        background_tasks.add_task(svc.run_single_phase, project_id, phase_id)
        return {"status": "phase_started", "phase_id": phase_id, "project_id": project_id}
    except ValueError as exc:
        raise HTTPException(400, str(exc))


class ResetPhasesRequest(BaseModel):
    phase_ids: List[str]

@app.post("/api/v1/projects/{project_id}/phases/reset", tags=["pipeline"])
async def reset_phases(project_id: int, body: ResetPhasesRequest, background_tasks: BackgroundTasks):
    """
    Reset given phases to 'pending' then immediately re-run the pipeline.
    Used by the frontend 'Re-run all stale' button after requirements change.
    """
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    if not body.phase_ids:
        raise HTTPException(400, "phase_ids must not be empty")

    # Validate phase IDs
    invalid = [p for p in body.phase_ids if p not in VALID_PHASES]
    if invalid:
        raise HTTPException(400, f"Invalid phase IDs: {invalid}")

    # Reset each phase status to pending
    svc = _project_svc()
    for phase_id in body.phase_ids:
        svc.set_phase_status(project_id, phase_id, "pending")

    log.info("api.phases_reset", extra={"project_id": project_id, "phase_ids": body.phase_ids})

    # Kick off the pipeline — it will now run all the reset phases
    pipeline = _pipeline_svc()
    background_tasks.add_task(pipeline.run_pipeline, project_id)
    return {"status": "pipeline_started", "reset_phases": body.phase_ids, "project_id": project_id}


@app.get("/api/v1/projects/{project_id}/export", tags=["projects"])
async def export_project_zip(project_id: int):
    """
    Stream all project output documents as a ZIP archive.
    Frontend uses this for the 'Download All Documents' button.
    """
    import io
    import zipfile
    import pathlib

    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "No output directory for this project")

    out_path = pathlib.Path(output_dir)
    if not out_path.exists():
        raise HTTPException(404, "Output directory does not exist")

    doc_files = [f for f in out_path.rglob("*") if f.is_file()]
    if not doc_files:
        raise HTTPException(404, "No documents found to export")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in doc_files:
            zf.write(f, f.relative_to(out_path))
    buf.seek(0)

    safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in (proj.get("name") or "project"))
    filename = f"{safe_name}_documents.zip"

    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _render_mermaid_diagrams_sync(md_text: str, tmp_dir: str) -> str:
    """
    Pre-render ```mermaid``` blocks to PNG via mermaid.ink API.
    All diagrams are fetched IN PARALLEL (ThreadPoolExecutor) with a 3s timeout each,
    so a document with N diagrams takes ~3 s total, not N×3 s.
    Failures fall back gracefully to a labelled code block.
    """
    import re as _re
    import base64 as _b64
    import urllib.request as _urlreq
    import pathlib as _pl
    from concurrent.futures import ThreadPoolExecutor, as_completed

    MERMAID_RE = _re.compile(r'```mermaid\s*\n([\s\S]*?)```', _re.IGNORECASE)
    tmp = _pl.Path(tmp_dir)

    # ── 1. Collect all mermaid blocks ─────────────────────────────────────────
    blocks = []  # list of (match, code)
    for m in MERMAID_RE.finditer(md_text):
        blocks.append((m, m.group(1).strip()))

    if not blocks:
        return md_text

    # ── 2. Fetch all diagrams in parallel ─────────────────────────────────────
    def fetch_diagram(idx_code):
        idx, code = idx_code
        try:
            encoded = _b64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
            url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"
            req = _urlreq.Request(url, headers={"User-Agent": "HardwarePipeline/1.0"})
            with _urlreq.urlopen(req, timeout=3) as resp:  # 3s max
                data = resp.read()
            if data and len(data) > 200:
                img_path = tmp / f"diagram_{idx}.png"
                img_path.write_bytes(data)
                return idx, str(img_path)
        except Exception as e:
            log.debug("mermaid.ink.skip idx=%d: %s", idx, e)
        return idx, None

    results: dict[int, str | None] = {}
    with ThreadPoolExecutor(max_workers=min(len(blocks), 6)) as pool:
        futures = {pool.submit(fetch_diagram, (i + 1, code)): i for i, (_, code) in enumerate(blocks)}
        for fut in as_completed(futures):
            idx, path = fut.result()
            results[idx] = path

    # ── 3. Replace blocks in reverse order (preserves string offsets) ─────────
    result_md = md_text
    for i, (m, code) in reversed(list(enumerate(blocks))):
        idx = i + 1
        img_path = results.get(idx)
        if img_path:
            replacement = f"\n\n**System Architecture Diagram {idx}**\n\n![Diagram {idx}]({img_path})\n\n"
        else:
            replacement = (
                f"\n\n**System Architecture Diagram {idx}** "
                f"*(rendered in browser — source below)*\n\n"
                f"```\n{code}\n```\n\n"
            )
        result_md = result_md[:m.start()] + replacement + result_md[m.end():]

    return result_md


@app.get("/api/v1/projects/{project_id}/docx/{filename:path}", tags=["projects"])
async def convert_document_to_docx(project_id: int, filename: str):
    """
    Convert a Markdown (.md) file to .docx and stream it for download.
    Mermaid diagrams are pre-rendered to PNG via mermaid.ink before conversion.
    Uses pandoc if available, falls back to python-docx.
    """
    import subprocess
    import tempfile
    import pathlib

    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "Project output directory not found")

    src_path = pathlib.Path(output_dir) / filename
    if not src_path.exists():
        raise HTTPException(404, f"File {filename} not found")

    if src_path.suffix.lower() not in (".md", ".txt"):
        raise HTTPException(400, "Only .md and .txt files can be converted to .docx")

    stem = src_path.stem
    out_filename = f"{stem}.docx"

    # ── Try pandoc first (installed in Docker image) ───────────────────────────
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Pre-render Mermaid blocks → PNG images in a thread (non-blocking)
            import asyncio
            import functools
            raw_md = src_path.read_text(encoding="utf-8")
            loop = asyncio.get_event_loop()
            processed_md = await loop.run_in_executor(
                None, functools.partial(_render_mermaid_diagrams_sync, raw_md, tmpdir)
            )

            # Write the processed markdown to a temp file in the same tmpdir
            # (so relative image paths resolve correctly for pandoc)
            tmp_md = pathlib.Path(tmpdir) / f"{stem}_processed.md"
            tmp_md.write_text(processed_md, encoding="utf-8")

            out_path = pathlib.Path(tmpdir) / out_filename
            result = subprocess.run(
                ["pandoc", str(tmp_md), "-o", str(out_path),
                 "--from=markdown", "--to=docx",
                 "-V", "geometry:margin=2.5cm",
                 "--standalone"],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0 and out_path.exists():
                data = out_path.read_bytes()
                return StreamingResponse(
                    iter([data]),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
                )
            log.warning("pandoc.failed", extra={"stderr": result.stderr, "file": filename})
    except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
        log.warning("pandoc.unavailable", extra={"error": str(exc)})

    # ── Fallback: python-docx heading/paragraph parser ────────────────────────
    try:
        from docx import Document as DocxDocument  # type: ignore

        md_text = src_path.read_text(encoding="utf-8")
        doc = DocxDocument()
        for line in md_text.splitlines():
            s = line.strip()
            if s.startswith("#### "):
                doc.add_heading(s[5:], level=4)
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s:
                doc.add_paragraph(s)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            data = pathlib.Path(tmp.name).read_bytes()

        return StreamingResponse(
            iter([data]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
        )
    except ImportError:
        raise HTTPException(500, "pandoc not found and python-docx not installed.")
    except Exception as exc:
        log.exception("docx.conversion_failed", extra={"file": filename})
        raise HTTPException(500, f"Conversion failed: {exc}")


# ── Phase status (polling endpoint for UI) ─────────────────────────────────────

@app.get("/api/v1/projects/{project_id}/status", tags=["pipeline"])
async def get_project_status(project_id: int):
    """Lightweight status poll — returns phase_statuses without full conversation."""
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    return {
        "project_id": project_id,
        "current_phase": proj.get("current_phase"),
        "phase_statuses": proj.get("phase_statuses", {}),
    }


# ── Test UI (standalone HTML workflow tester) ──────────────────────────────────

@app.get("/testui", response_class=HTMLResponse, tags=["ops"])
async def test_ui():
    """Serve standalone HTML workflow test page (no Streamlit needed)."""
    import pathlib
    p = pathlib.Path(__file__).parent / "test_ui.html"
    if p.exists():
        return HTMLResponse(content=p.read_text(), status_code=200)
    return HTMLResponse(content="<h1>test_ui.html not found</h1>", status_code=404)


@app.get("/app", response_class=HTMLResponse, tags=["ops"])
async def serve_frontend():
    """Serve the React v5 frontend bundle at http://localhost:8000/app"""
    import pathlib
    p = pathlib.Path(__file__).parent / "frontend" / "bundle.html"
    if p.exists():
        return HTMLResponse(content=p.read_text(encoding="utf-8", errors="replace"), status_code=200)
    return HTMLResponse(content="<h1>Frontend not built yet. Run the React build.</h1>", status_code=404)


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host=settings.fastapi_host,
        port=settings.fastapi_port,
        reload=settings.debug,
        log_level=settings.log_level.lower(),
    )
